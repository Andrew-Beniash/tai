"""
Document parsing utilities for extracting text from various file formats.
Supports PDF, DOCX, XLSX, and text files.
Integrates with Google Drive for document retrieval.
"""

import logging
import os
import re
from typing import Optional, Dict, Any, List
import io
import base64
from pathlib import Path

# For PDF parsing
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# For DOCX parsing
try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# For XLSX parsing
try:
    import pandas as pd
    XLSX_SUPPORT = True
except ImportError:
    XLSX_SUPPORT = False

from app.core.drive_client import drive_client
from app.services.document_service import document_service

# Configure logging
logger = logging.getLogger(__name__)

# Dummy content for prototype - in a real app, these would be parsed from actual files
DUMMY_DOCUMENTS = {
    "prior_year_return.pdf": """
ACME CORPORATION - FORM 1120
Tax Year 2023
EIN: 12-3456789

INCOME:
Total Revenue: $5,435,000
Cost of Goods: $2,150,000
Gross Profit: $3,285,000
Operating Expenses: $2,100,000
Net Income: $1,185,000

TAX CALCULATION:
Taxable Income: $1,100,000
Federal Tax Rate: 21%
Federal Tax: $231,000

NOTES:
- Depreciation method for new equipment to be reviewed
- Potential R&D credit application for software development
- Missing documentation for charitable contributions
- Foreign income from Canadian subsidiary requires additional forms
""",
    "financial_statement.xlsx": """
ACME CORPORATION
Balance Sheet as of Dec 31, 2023

ASSETS:
Current Assets: $2,750,000
Fixed Assets: $4,200,000
Total Assets: $6,950,000

LIABILITIES:
Current Liabilities: $1,250,000
Long-term Debt: $2,340,000
Total Liabilities: $3,590,000

EQUITY:
Common Stock: $1,000,000
Retained Earnings: $2,360,000
Total Equity: $3,360,000

INCOME STATEMENT:
Revenue: $5,435,000
Expenses: $4,250,000
Net Income: $1,185,000

CASH FLOW:
Operating Activities: $1,230,000
Investing Activities: ($850,000)
Financing Activities: ($300,000)
Net Change in Cash: $80,000
""",
    "SOW.docx": """
STATEMENT OF WORK
Client: ACME Corporation
Tax Year: 2024
Services: Corporate Tax Filing (Form 1120)

SCOPE:
- Preparation of Form 1120 and all required schedules
- Tax planning advisory services
- Quarterly estimated tax payment calculations
- State tax returns for CA, NY, TX

TIMELINE:
- Initial documentation due: February 28, 2024
- Draft return review: March 15, 2024
- Final filing deadline: April 15, 2024

FEES:
Base preparation fee: $12,500
Additional services billed at $250/hour

NOTES:
- Client has expanded operations to Canada requiring international tax considerations
- New manufacturing facility may qualify for additional deductions
- CEO compensation package requires special documentation
""",
    "client_responses.docx": """
ACME CORPORATION
Responses to Tax Questionnaire

1. Has the company structure changed? YES
   Details: Added Canadian subsidiary in June 2023

2. Any new major assets purchased? YES
   Details: Manufacturing equipment ($1.2M) in August 2023

3. Any changes to officer compensation? YES
   Details: New CEO package includes stock options

4. Any new debt or financing? NO

5. Any legal settlements or lawsuits? NO

6. Any foreign operations or accounts? YES
   Details: Canadian operations began June 2023
   
7. Any tax credits being claimed? UNSURE
   Details: May qualify for R&D credits for software development

MISSING ITEMS:
- Detailed breakdown of R&D expenses
- Officer compensation documentation
- Final depreciation schedules
- Foreign income statements
""",
    "form_1120_template.docx": """
FORM 1120 - U.S. Corporation Income Tax Return
Tax Year: 2024

Part I: Income
1. Gross receipts or sales: _______
2. Returns and allowances: _______
3. Cost of goods sold: _______
4. Gross profit (subtract line 3 from line 1c): _______
5. Dividends: _______
6. Interest: _______
7. Gross rents: _______
8. Gross royalties: _______
9. Capital gain net income: _______
10. Net gain or (loss) from Form 4797: _______

Part II: Deductions
12. Compensation of officers: _______
13. Salaries and wages: _______
14. Repairs and maintenance: _______
15. Bad debts: _______
16. Rents: _______
17. Taxes and licenses: _______
18. Interest: _______
19. Charitable contributions: _______
20. Depreciation: _______
21. Depletion: _______
22. Advertising: _______
23. Pension, profit-sharing plans: _______

Part III: Tax Computation
31. Taxable income: _______
32. Total tax: _______
""",
    "form_1065_template.docx": """
FORM 1065 - U.S. Return of Partnership Income
Tax Year: 2024

Part I: Income
1. Gross receipts or sales: _______
2. Returns and allowances: _______
3. Cost of goods sold: _______
4. Gross profit (subtract line 3 from line 1c): _______
5. Ordinary income (loss) from other partnerships: _______
6. Net farm profit (loss): _______
7. Net gain (loss) from Form 4797: _______
8. Other income (loss): _______

Part II: Deductions
9. Salaries and wages: _______
10. Guaranteed payments to partners: _______
11. Repairs and maintenance: _______
12. Bad debts: _______
13. Rent: _______
14. Taxes and licenses: _______
15. Interest: _______
16. Depreciation: _______
17. Depletion: _______
18. Retirement plans: _______
19. Employee benefit programs: _______
20. Other deductions: _______

Schedule K: Partners' Distributive Share Items
1. Ordinary business income (loss): _______
2. Net rental real estate income (loss): _______
3. Other net rental income (loss): _______
4. Guaranteed payments: _______
5. Interest income: _______
6. Dividends: _______
7. Royalties: _______
8. Net short-term capital gain (loss): _______
9. Net long-term capital gain (loss): _______
"""
}

async def extract_document_text(doc_id: str, filename: str = None) -> str:
    """
    Extract text content from a document.
    Uses document_service to fetch the document content from Google Drive,
    then extracts text based on file type.
    
    Args:
        doc_id: The document ID
        filename: Optional filename (if doc_id is a Google Drive file ID rather than our internal ID)
        
    Returns:
        Extracted text content
    """
    logger.info(f"Extracting text from document {doc_id}")
    
    try:
        # For prototype, check if we have dummy content
        if filename and filename in DUMMY_DOCUMENTS:
            logger.info(f"Using dummy content for {filename}")
            return DUMMY_DOCUMENTS[filename]
        
        # First try to get the document from our document service
        if not filename:
            # Try to get document metadata from our database
            document = await document_service.get_by_id(doc_id)
            if document:
                filename = document.file_name
        
        # Get document content from Google Drive via our document service
        text_content = await document_service.get_text_content(doc_id)
        if text_content:
            return text_content
        
        # If text content extraction failed or isn't implemented, get raw content and parse it
        content, mime_type = await document_service.get_document_content(doc_id)
        if not content:
            logger.warning(f"Could not retrieve document content for {doc_id}")
            return f"[Could not retrieve document content]"
        
        # Extract text based on file type
        if not filename:
            # If no filename, try to determine from mime_type
            if mime_type == 'application/pdf':
                return extract_text_from_pdf_bytes(content)
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return extract_text_from_docx_bytes(content)
            elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                return extract_text_from_xlsx_bytes(content)
            elif mime_type and mime_type.startswith('text/'):
                return content.decode('utf-8', errors='replace')
            else:
                return f"[Content extraction not supported for this file type: {mime_type}]"
        else:
            # If we have a filename, use the extension
            _, ext = os.path.splitext(filename.lower())
            
            if ext == '.pdf':
                return extract_text_from_pdf_bytes(content)
            elif ext == '.docx':
                return extract_text_from_docx_bytes(content)
            elif ext == '.xlsx':
                return extract_text_from_xlsx_bytes(content)
            elif ext in ['.txt', '.csv', '.json', '.md']:
                return content.decode('utf-8', errors='replace')
            else:
                logger.warning(f"Unsupported file type: {ext} for {filename}")
                return f"[Content extraction not supported for {ext} files]"
            
    except Exception as e:
        logger.error(f"Error extracting text from document {doc_id}: {str(e)}")
        return f"[Error extracting content: {str(e)}]"

def extract_text_from_pdf_bytes(content: bytes) -> str:
    """
    Extract text from PDF content.
    
    Args:
        content: PDF content as bytes
        
    Returns:
        Extracted text
    """
    if not PDF_SUPPORT:
        return "[PDF parsing support not installed. Install PyPDF2 package.]"
    
    try:
        # Create a PDF reader object
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        # Extract text from each page
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n\n"
        
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        return f"[Error extracting PDF text: {str(e)}]"

def extract_text_from_docx_bytes(content: bytes) -> str:
    """
    Extract text from DOCX content.
    
    Args:
        content: DOCX content as bytes
        
    Returns:
        Extracted text
    """
    if not DOCX_SUPPORT:
        return "[DOCX parsing support not installed. Install python-docx package.]"
    
    try:
        # Create a DOCX document
        docx_file = io.BytesIO(content)
        doc = docx.Document(docx_file)
        
        # Extract text from paragraphs
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        return f"[Error extracting DOCX text: {str(e)}]"

def extract_text_from_xlsx_bytes(content: bytes) -> str:
    """
    Extract text from XLSX content.
    
    Args:
        content: XLSX content as bytes
        
    Returns:
        Extracted text
    """
    if not XLSX_SUPPORT:
        return "[XLSX parsing support not installed. Install pandas package.]"
    
    try:
        # Create an XLSX file
        xlsx_file = io.BytesIO(content)
        
        # Read all sheets
        xlsx_data = pd.read_excel(xlsx_file, sheet_name=None)
        
        # Extract text from each sheet
        text = ""
        for sheet_name, df in xlsx_data.items():
            text += f"SHEET: {sheet_name}\n"
            text += df.to_string() + "\n\n"
        
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from XLSX: {str(e)}")
        return f"[Error extracting XLSX text: {str(e)}]"

async def get_document_preview(doc_id: str, filename: str = None, max_length: int = 500) -> str:
    """
    Get a preview of a document's content.
    
    Args:
        doc_id: The document ID
        filename: Optional filename
        max_length: Maximum length of preview
        
    Returns:
        Document preview
    """
    try:
        # Get full document text
        full_text = await extract_document_text(doc_id, filename)
        
        # Trim to max length
        if len(full_text) > max_length:
            preview = full_text[:max_length] + "..."
        else:
            preview = full_text
            
        return preview
    
    except Exception as e:
        logger.error(f"Error getting document preview: {str(e)}")
        return f"[Error getting preview: {str(e)}]"

async def get_documents_content_for_task(task_id: str, max_chars_per_doc: int = 10000) -> Dict[str, str]:
    """
    Get text content for all documents associated with a task.
    
    Args:
        task_id: Task ID
        max_chars_per_doc: Maximum characters to extract per document
        
    Returns:
        Dictionary mapping document IDs to their text content
    """
    # Get documents for task
    documents = await document_service.get_documents_for_task(task_id)
    
    result = {}
    for doc in documents:
        # Extract text content
        text = await extract_document_text(doc.doc_id)
        
        # Trim if necessary
        if len(text) > max_chars_per_doc:
            text = text[:max_chars_per_doc] + "... [content truncated]"
            
        result[doc.doc_id] = text
    
    return result

async def get_document_metadata(project_id: str = None, task_id: str = None) -> List[Dict[str, Any]]:
    """
    Get metadata for documents associated with a project or task.
    
    Args:
        project_id: Optional project ID
        task_id: Optional task ID
        
    Returns:
        List of document metadata dictionaries
    """
    documents = []
    
    if project_id:
        # Get all documents for project
        documents = await document_service.get_documents_by_project(project_id)
    elif task_id:
        # Get documents for task
        documents = await document_service.get_documents_for_task(task_id)
    
    # Convert to dictionary format
    result = []
    for doc in documents:
        result.append({
            "doc_id": doc.doc_id,
            "file_name": doc.file_name,
            "file_type": doc.file_type,
            "last_modified": doc.last_modified.isoformat(),
            "size_bytes": doc.size_bytes,
            "web_view_link": doc.web_view_link
        })
    
    return result
